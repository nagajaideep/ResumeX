"""
Text extraction from PDF, DOCX, and TEX files.
"""

import io
import pdfplumber
import docx


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Extract plain text from a file based on its extension."""

    name = filename.lower()

    if name.endswith(".pdf"):
        return _extract_pdf(file_bytes)
    elif name.endswith(".docx"):
        return _extract_docx(file_bytes)
    elif name.endswith(".tex"):
        return file_bytes.decode("utf-8", errors="replace")
    else:
        raise ValueError(f"Unsupported file type: {filename}")


def _extract_pdf(data: bytes) -> str:
    """Extract text from all pages of a PDF."""
    pages = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
    return "\n\n".join(pages)


def _extract_docx(data: bytes) -> str:
    """Extract text from paragraphs and tables in a DOCX."""
    doc = docx.Document(io.BytesIO(data))
    parts = []

    # Paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)
